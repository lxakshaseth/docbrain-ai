import os
from typing import List
from pypdf import PdfReader

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.docstore.document import Document

from app.core.config import settings
from app.core.logger import logger
from app.core.exceptions import IngestionError

class PDFProcessor:
    @staticmethod
    def extract_and_chunk(file_path: str, chunk_size: int = None, chunk_overlap: int = None) -> List[Document]:
        if not os.path.exists(file_path):
            raise IngestionError(f"File does not exist at path: {file_path}")

        c_size = chunk_size or settings.default_chunk_size
        c_overlap = chunk_overlap or settings.default_chunk_overlap
        ext = os.path.splitext(file_path)[1].lower()

        try:
            logger.info(f"Loading document ({ext}): {file_path}")
            raw_docs = []

            if ext == ".pdf":
                # Primary PDF extraction via PyPDF
                try:
                    reader = PdfReader(file_path)
                    total_pages = len(reader.pages)
                    logger.info(f"Processing PDF with {total_pages} total pages...")

                    for i, page in enumerate(reader.pages):
                        page_num = i + 1
                        text = page.extract_text() or ""
                        
                        # Fallback for scanned page or diagram image on page i
                        if not text.strip():
                            try:
                                import fitz
                                fitz_doc = fitz.open(file_path)
                                if i < len(fitz_doc):
                                    text = fitz_doc[i].get_text() or ""
                            except Exception:
                                pass

                        if not text.strip():
                            text = f"[Document Page {page_num} contains visual diagrams, images, or formatted tables]"

                        raw_docs.append(Document(
                            page_content=f"--- Page {page_num} ---\n{text}",
                            metadata={"page": page_num, "source": file_path, "totalPages": total_pages}
                        ))
                except Exception as pe:
                    logger.warning(f"PyPDF reader failed, attempting PyMuPDF fallback: {pe}")

                # Secondary Fallback if PyPDF fails completely
                if not raw_docs:
                    try:
                        import fitz
                        doc = fitz.open(file_path)
                        total_pages = len(doc)
                        for i, page in enumerate(doc):
                            page_num = i + 1
                            text = page.get_text() or f"[Page {page_num} visual content]"
                            raw_docs.append(Document(
                                page_content=f"--- Page {page_num} ---\n{text}",
                                metadata={"page": page_num, "source": file_path, "totalPages": total_pages}
                            ))
                    except Exception as fe:
                        logger.warning(f"PyMuPDF fallback failed: {fe}")

            elif ext == ".docx":
                try:
                    import docx
                    doc = docx.Document(file_path)
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    full_text = "\n\n".join(paragraphs) if paragraphs else "[Docx document content]"
                    raw_docs.append(Document(
                        page_content=full_text,
                        metadata={"page": 1, "source": file_path, "totalPages": 1}
                    ))
                except Exception as de:
                    logger.error(f"Docx parsing error: {de}")

            elif ext in [".txt", ".md"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                    if not content.strip():
                        content = "[Text document content]"
                    raw_docs.append(Document(
                        page_content=content,
                        metadata={"page": 1, "source": file_path, "totalPages": 1}
                    ))

            if not raw_docs:
                raise IngestionError(f"No readable content found in document: {file_path}")

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=c_size,
                chunk_overlap=c_overlap,
                separators=["\n--- Page", "\n\n", "\n", " ", ""]
            )

            chunks = text_splitter.split_documents(raw_docs)
            logger.info(f"Extracted {len(chunks)} chunks across {len(raw_docs)} pages from {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error extracting text from document: {e}")
            raise IngestionError(f"Failed to process document: {str(e)}")
